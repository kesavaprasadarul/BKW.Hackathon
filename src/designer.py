"""Report Designer"""
from datetime import datetime
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import darkblue
from docx import Document
from config import REPORTS_DIR


class Designer:
    def __init__(self):
        REPORTS_DIR.mkdir(exist_ok=True)
    
    def _filename(self, ext):
        return f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
    
    def pdf(self, content):
        path = REPORTS_DIR / self._filename('.pdf')
        doc = SimpleDocTemplate(str(path), pagesize=A4)
        styles = getSampleStyleSheet()
        
        title = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=16, 
                              spaceAfter=20, textColor=darkblue, alignment=1)
        heading = ParagraphStyle('Heading', parent=styles['Heading2'], 
                                fontSize=12, spaceAfter=10, textColor=darkblue)
        body = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, spaceAfter=8)
        
        story = [
            Paragraph("AI Generated Report", title),
            Spacer(1, 0.2*inch),
            Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']),
            Spacer(1, 0.3*inch)
        ]
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1*inch))
            elif line.startswith('#') or (len(line) < 50 and line.isupper()):
                text = line.lstrip('#').strip()
                if text:
                    story.append(Paragraph(text, heading))
            else:
                story.append(Paragraph(line, body))
        
        doc.build(story)
        return str(path)
    
    def docx(self, content):
        path = REPORTS_DIR / self._filename('.docx')
        doc = Document()
        doc.add_heading('AI Generated Report', 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('#'):
                level = min(len(line) - len(line.lstrip('#')), 3)
                doc.add_heading(line.lstrip('#').strip(), level)
            elif len(line) < 50 and line.isupper():
                doc.add_heading(line, 2)
            else:
                doc.add_paragraph(line)
        
        doc.save(str(path))
        return str(path)
    
    def markdown(self, content):
        path = REPORTS_DIR / self._filename('.md')
        md = f"# AI Generated Report\n\n**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n---\n\n{content}\n"
        path.write_text(md, encoding='utf-8')
        return str(path)
