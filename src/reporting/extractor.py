"""File Extractor for Multiple Formats"""
import pandas as pd
import PyPDF2
import docx
from pathlib import Path
import logging
from typing import Dict, List, Optional, Union
import re
import warnings

from pandas.io import parsers

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileExtractor:
    """Extracts text content from various file formats"""
    
    def __init__(self):
        self.supported_extensions = {
            '.txt': self._extract_txt,
            '.md': self._extract_txt,
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.xlsx': self._extract_excel,
            '.xls': self._extract_excel,
            '.csv': self._extract_csv
        }
    
    def extract_from_directory(self, directory_path: Union[str, Path]) -> Dict[str, str]:
        """
        Extract text content from all supported files in a directory
        
        Args:
            directory_path: Path to directory containing files
            
        Returns:
            Dictionary mapping filename to extracted content
        """
        directory = Path(directory_path)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        extracted_data = {}
        
        # Get all files in directory and subdirectories
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    content = self.extract_from_file(file_path)
                    if content:
                        extracted_data[file_path.name] = content
                        logger.info(f"Successfully extracted content from: {file_path.name}")
                except Exception as e:
                    logger.error(f"Failed to extract from {file_path.name}: {e}")
        
        return extracted_data
    
    def extract_from_file(self, file_path: Union[str, Path]) -> Optional[str]:
        """
        Extract text content from a single file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content or None if extraction fails
        """
        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        extension = file_path.suffix.lower()
        if extension not in self.supported_extensions:
            logger.warning(f"Unsupported file type: {extension}")
            return None
        
        try:
            extractor_func = self.supported_extensions[extension]
            return extractor_func(file_path)
        except Exception as e:
            logger.error(f"Error extracting from {file_path.name}: {e}")
            return None
    
    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT and MD files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    return file_path.read_text(encoding=encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode file with any supported encoding")
        except Exception as e:
            logger.error(f"Error reading text file {file_path.name}: {e}")
            return ""
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"--- Seite {page_num + 1} ---\n{text}")
            
            return "\n\n".join(text_content)
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path.name}: {e}")
            return ""
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    paragraphs.append("Tabelle:\n" + "\n".join(table_text))
            
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path.name}: {e}")
            return ""
    
    def _extract_excel(self, file_path: Path) -> str:
        """Extract text from Excel files"""
        try:
            excel_content = []
            
            # suppress warnings (openpyxl)
            with warnings.catch_warnings():
                warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')
                
                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Add sheet header
                    excel_content.append(f"=== Arbeitsblatt: {sheet_name} ===")
                    
                    # Convert DataFrame to readable text
                    if not df.empty:
                        # Get column names
                        columns = df.columns.tolist()
                        excel_content.append(f"Spalten: {', '.join(map(str, columns))}")
                        
                        # Add data rows (limit to first 100 rows to avoid too much content)
                        max_rows = min(100, len(df))
                        for i in range(max_rows):
                            row_data = []
                            for col in columns:
                                value = df.iloc[i][col]
                                if pd.notna(value):
                                    row_data.append(str(value))
                                else:
                                    row_data.append("")
                            excel_content.append(" | ".join(row_data))
                        
                        if len(df) > 100:
                            excel_content.append(f"... und {len(df) - 100} weitere Zeilen")
                    
                    excel_content.append("")  # Empty line between sheets
            
            return "\n".join(excel_content)
        except Exception as e:
            logger.error(f"Error reading Excel file {file_path.name}: {e}")
            return ""
    
    def _extract_csv(self, file_path: Path) -> str:
        """Extract text from CSV files"""
        try:
            df = pd.read_csv(file_path)
            
            csv_content = []
            csv_content.append(f"=== CSV Datei: {file_path.name} ===")
            
            if not df.empty:
                # Get column names
                columns = df.columns.tolist()
                csv_content.append(f"Spalten: {', '.join(map(str, columns))}")
                
                # Add data rows (limit to first 100 rows)
                max_rows = min(100, len(df))
                for i in range(max_rows):
                    row_data = []
                    for col in columns:
                        value = df.iloc[i][col]
                        if pd.notna(value):
                            row_data.append(str(value))
                        else:
                            row_data.append("")
                    csv_content.append(" | ".join(row_data))
                
                if len(df) > 100:
                    csv_content.append(f"... und {len(df) - 100} weitere Zeilen")
            
            return "\n".join(csv_content)
        except Exception as e:
            logger.error(f"Error reading CSV file {file_path.name}: {e}")
            return ""
    
    def combine_extracted_data(self, extracted_data: Dict[str, str], 
                             project_name: str = "Projekt") -> str:
        """
        Combine extracted data from multiple files into a structured format
        
        Args:
            extracted_data: Dictionary mapping filename to content
            project_name: Name of the project
            
        Returns:
            Combined text content suitable for report generation
        """
        combined_content = []
        combined_content.append(f"# Projektdaten - {project_name}")
        combined_content.append(f"Extrahierte Informationen aus {len(extracted_data)} Dateien")
        combined_content.append("=" * 50)
        combined_content.append("")
        
        # Sort files by type and name for better organization
        file_categories = {
            'Berichte': [],
            'Kosten': [],
            'Berechnungen': [],
            'Pläne': [],
            'Sonstige': []
        }
        
        for filename, content in extracted_data.items():
            filename_lower = filename.lower()
            if any(keyword in filename_lower for keyword in ['bericht', 'report', 'erläuterung']):
                file_categories['Berichte'].append((filename, content))
            elif any(keyword in filename_lower for keyword in ['kosten', 'cost', 'preis', 'kobe', 'kosch']):
                file_categories['Kosten'].append((filename, content))
            elif any(keyword in filename_lower for keyword in ['berechnung', 'calculation', 'heizlast', 'kühllast']):
                file_categories['Berechnungen'].append((filename, content))
            elif any(keyword in filename_lower for keyword in ['plan', 'schema', 'grundriss']):
                file_categories['Pläne'].append((filename, content))
            else:
                file_categories['Sonstige'].append((filename, content))
        
        # Add content by category
        for category, files in file_categories.items():
            if files:
                combined_content.append(f"## {category}")
                combined_content.append("")
                
                for filename, content in files:
                    combined_content.append(f"### {filename}")
                    combined_content.append("")
                    
                    # Clean and format content
                    cleaned_content = self._clean_content(content)
                    combined_content.append(cleaned_content)
                    combined_content.append("")
                    combined_content.append("---")
                    combined_content.append("")
        
        return "\n".join(combined_content)
    
    def _clean_content(self, content: str) -> str:
        """Clean and format extracted content"""
        if not content:
            return ""
        
        # Remove excessive whitespace
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        # Remove very short lines that might be artifacts
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if len(line) > 2 or line in ['', '---', '===']:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)


def _extract_detailed_room_data(df: pd.DataFrame, col_mapping: Dict[str, str]) -> List[Dict]:
    """Extract detailed room-by-room data from Excel DataFrame"""
    detailed_rooms = []
    
    extended_mapping = {}
    for col in df.columns:
        col_str = str(col).lower()
        
        if 'gebäudeteil' in col_str:
            extended_mapping['gebäudeteil'] = col
        elif 'geschoss' in col_str:
            extended_mapping['geschoss'] = col
        elif 'raum-nr' in col_str and 'zusätzliche' not in col_str:
            extended_mapping['raum_nr'] = col
        elif 'raum-bezeichnung' in col_str:
            extended_mapping['raum_bezeichnung'] = col
        elif 'nummer' in col_str and 'raumtyp' in col_str:
            extended_mapping['raumtyp_nr'] = col
        elif 'bezeichnung' in col_str and 'raumtyp' in col_str:
            extended_mapping['raumtyp_bez'] = col
        elif 'fläche' in col_str:
            extended_mapping['fläche'] = col
        elif 'volumen' in col_str:
            extended_mapping['volumen'] = col
        elif 'höhe' in col_str and 'mep' in col_str:
            extended_mapping['höhe_mep'] = col
        elif 'lichte' in col_str and 'raumhöhe' in col_str:
            extended_mapping['lichte_raumhöhe'] = col
        elif 'abhang' in col_str and 'decke' in col_str and 'typ' not in col_str and 'höhe' not in col_str:
            extended_mapping['abhangdecke'] = col
        elif 'typ' in col_str and 'abhangdecke' in col_str:
            extended_mapping['abhangdecke_typ'] = col
        elif 'lichte' in col_str and 'höhe' in col_str and 'abhang' in col_str:
            extended_mapping['lichte_höhe_abhang'] = col
        elif 'temperatur' in col_str and 'winter' in col_str:
            extended_mapping['temp_winter'] = col
        elif 'temperatur' in col_str and 'sommer' in col_str:
            extended_mapping['temp_sommer'] = col
        elif 'beheizt' in col_str:
            extended_mapping['beheizt'] = col
        elif 'lüftung' in col_str and 'mech' in col_str:
            extended_mapping['mech_lüftung'] = col
        elif 'luftfeuchte' in col_str:
            extended_mapping['luftfeuchte'] = col
        elif 'akustik' in col_str:
            extended_mapping['akustik'] = col
        elif 'sprinkler' in col_str:
            extended_mapping['sprinkler'] = col
        elif 'anzahl' in col_str and 'personen' in col_str:
            extended_mapping['personen_anzahl'] = col
        elif 'kommentar' in col_str and '(1)' in col_str:
            extended_mapping['kommentar_1'] = col
        elif 'kommentar' in col_str and '(2)' in col_str:
            extended_mapping['kommentar_2'] = col
    
    numeric_fields = {'fläche', 'volumen', 'höhe_mep', 'lichte_raumhöhe', 'lichte_höhe_abhang', 'raumtyp_nr'}
    temp_fields = {'temp_winter', 'temp_sommer'}
    
    for idx, row in df.iterrows():
        if 'Teilergebnis' in str(row.values) or 'Summe' in str(row.values):
            continue
        
        raum_nr_col = extended_mapping.get('raum_nr', df.columns[9])
        raum_nr = str(row[raum_nr_col]) if raum_nr_col in df.columns else ''
        if raum_nr in ['nan', '', 'Raum-Nr.']:
            continue
        
        room_data = {}
        
        for field, col_name in extended_mapping.items():
            if col_name not in df.columns:
                continue
            
            value = row[col_name]
            if pd.isna(value):
                continue
            
            # Handle numeric fields
            if field in numeric_fields:
                try:
                    value = float(value)
                    if value > 0:
                        room_data[field] = value
                except:
                    pass
            # Handle temperature fields
            elif field in temp_fields:
                try:
                    value = float(value)
                    if value != 0:
                        room_data[field] = int(value) if value == int(value) else value
                except:
                    parsers
            else:
                value_str = str(value).strip()
                if value_str and value_str not in ['nan', '0']:
                    room_data[field] = value_str
        
        if len(room_data) > 2:
            detailed_rooms.append(room_data)
    
    return detailed_rooms


def extract_hvac_project_data(excel_path: Union[str, Path], json_path: Union[str, Path]) -> str:
    """
    Extract and structure data from Excel room data and JSON cost estimation
    
    Args:
        excel_path: Path to Excel file with room data
        json_path: Path to JSON file with BOQ and cost data
        
    Returns:
        Structured project data as text suitable for report generation
    """
    import json
    
    excel_path = Path(excel_path)
    json_path = Path(json_path)
    
    if not excel_path.exists() or not json_path.exists():
        logger.error("Input files not found")
        return ""
    
    content = []
    content.append("# PROJEKTDATEN - TGA-PLANUNG")
    content.append("=" * 80)
    content.append("")
    
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        content.append("## PROJEKTKENNZAHLEN")
        content.append("")
        
        if 'summary' in json_data and 'project_metrics' in json_data['summary']:
            metrics = json_data['summary']['project_metrics']
            content.append("**Technische Kennwerte:**")
            content.append(f"* Gesamtfläche: {metrics.get('Total Area (m^2)', 0):.2f} m²")
            content.append(f"* Gesamtvolumen: {metrics.get('Total Volume (m^3)', 0):.2f} m³")
            content.append(f"* Heizlast gesamt: {metrics.get('Total Heating Load (kW)', 0):.2f} kW")
            content.append(f"* Kühllast gesamt: {metrics.get('Total Cooling Load (kW)', 0):.2f} kW")
            content.append(f"* Luftvolumenstrom gesamt: {metrics.get('Total Airflow (m3/h)', 0):,.0f} m³/h")
            content.append("")
        
        if 'summary' in json_data and 'grand_total_cost' in json_data['summary']:
            content.append(f"**Gesamtkosten (geschätzt):** {json_data['summary']['grand_total_cost']:,.2f} EUR")
            content.append("")
        
        if 'summary' in json_data and 'cost_factors_applied' in json_data['summary']:
            factors = json_data['summary']['cost_factors_applied']
            content.append("**Angewendete Kostenfaktoren:**")
            content.append(f"* Regionalfaktor München: {factors.get('regional_factor_munich', 1.0)}")
            content.append(f"* Zeitindex: {factors.get('time_index_factor', 1.0)}")
            content.append(f"* Lohnfaktor: {factors.get('labor_factor', 1.0)}")
            content.append(f"* Gemeinkosten/Gewinn: {factors.get('overhead_profit_factor', 0.0)}")
            content.append("")
        
        # Extract BOQ by cost groups
        if 'detailed_boq' in json_data:
            content.append("## LEISTUNGSVERZEICHNIS NACH KOSTENGRUPPEN")
            content.append("")
            
            # Group items by KG
            kg_groups = {}
            for item in json_data['detailed_boq']:
                kg = item.get('subgroup_kg', 'Sonstige')
                kg_title = item.get('subgroup_title', 'Sonstige')
                key = f"KG {kg} - {kg_title}"
                
                if key not in kg_groups:
                    kg_groups[key] = []
                kg_groups[key].append(item)
            
            # Output by KG
            for kg_key in sorted(kg_groups.keys()):
                content.append(f"### {kg_key}")
                content.append("")
                
                items = kg_groups[kg_key]
                total_cost = sum(item.get('total_final_price', 0) for item in items)
                
                content.append(f"**Gesamtkosten {kg_key.split('-')[0].strip()}:** {total_cost:,.2f} EUR")
                content.append("")
                content.append("**Positionen:**")
                
                for item in items:
                    qty = item.get('quantity', 0)
                    unit = item.get('unit', '')
                    desc = item.get('description', '')
                    price = item.get('total_final_price', 0)
                    bki_ref = item.get('bki_component_title', '')
                    
                    # Show items with quantity or price (skip only if both are zero)
                    if qty > 0 or price > 0:
                        if price > 0:
                            content.append(f"* {desc}: {qty:.1f} {unit} - {price:,.2f} EUR")
                        else:
                            content.append(f"* {desc}: {qty:.1f} {unit} - nicht kalkuliert")
                        
                        if bki_ref and bki_ref != 'N/A' and 'Percentage Based' not in bki_ref:
                            content.append(f"  (BKI-Referenz: {bki_ref})")
                
                content.append("")
        
        logger.info("Successfully extracted JSON cost data")
    except Exception as e:
        logger.error(f"Error extracting JSON data: {e}")
    
    try:
        with warnings.catch_warnings():
            warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')
            
            df = pd.read_excel(excel_path, sheet_name=0)
            
            # Find data start row (after headers with "Raum-Nr.", "Raum-Bezeichnung", etc.)
            data_start_row = None
            for idx, row in df.iterrows():
                if 'Raum-Nr.' in str(row.values) or 'Raum-Bezeichnung' in str(row.values):
                    data_start_row = idx + 1
                    break
            
            if data_start_row is None:
                data_start_row = 7  # Default fallback
            
            # Re-read with proper header
            df = pd.read_excel(excel_path, sheet_name=0, header=data_start_row)
            
            # Clean up - remove completely empty rows
            df = df.dropna(how='all')
            
            # Filter out summary rows
            df = df[~df.iloc[:, 5].astype(str).str.contains('Teilergebnis|Summe', case=False, na=False)]
            
            content.append("## RAUMDATEN")
            content.append("")
            
            # Find relevant columns (flexible column detection)
            col_mapping = {}
            for col in df.columns:
                col_str = str(col).lower()
                if 'raum-nr' in col_str or 'raumnr' in col_str:
                    col_mapping['room_nr'] = col
                elif 'raum-bezeichnung' in col_str or 'raumbezeichnung' in col_str:
                    col_mapping['room_name'] = col
                elif 'raumtyp' in col_str and 'nummer' in col_str:
                    col_mapping['room_type_nr'] = col
                elif 'raumtyp' in col_str and 'bezeichnung' in col_str:
                    col_mapping['room_type_name'] = col
                elif 'fläche' in col_str:
                    col_mapping['area'] = col
                elif 'volumen' in col_str:
                    col_mapping['volume'] = col
                elif 'temperatur' in col_str and 'winter' in col_str:
                    col_mapping['temp_winter'] = col
                elif 'temperatur' in col_str and 'sommer' in col_str:
                    col_mapping['temp_summer'] = col
                elif 'beheizt' in col_str:
                    col_mapping['heated'] = col
                elif 'lüftung' in col_str:
                    col_mapping['ventilation'] = col
            
            # Group rooms by type
            room_types = {}
            total_area = 0
            total_volume = 0
            
            for idx, row in df.iterrows():
                try:
                    room_type = str(row[col_mapping.get('room_type_name', df.columns[11])]) if 'room_type_name' in col_mapping else 'Unbekannt'
                    
                    if room_type == 'nan' or room_type == 'Unbekannt':
                        continue
                    
                    if room_type not in room_types:
                        room_types[room_type] = {
                            'count': 0,
                            'total_area': 0,
                            'total_volume': 0,
                            'rooms': []
                        }
                    
                    room_nr = str(row[col_mapping.get('room_nr', df.columns[9])]) if 'room_nr' in col_mapping else ''
                    room_name = str(row[col_mapping.get('room_name', df.columns[10])]) if 'room_name' in col_mapping else ''
                    area = float(row[col_mapping.get('area', df.columns[13])]) if 'area' in col_mapping and pd.notna(row[col_mapping.get('area', df.columns[13])]) else 0
                    volume = float(row[col_mapping.get('volume', df.columns[14])]) if 'volume' in col_mapping and pd.notna(row[col_mapping.get('volume', df.columns[14])]) else 0
                    
                    if area > 0:
                        room_types[room_type]['count'] += 1
                        room_types[room_type]['total_area'] += area
                        room_types[room_type]['total_volume'] += volume
                        room_types[room_type]['rooms'].append({
                            'nr': room_nr,
                            'name': room_name,
                            'area': area,
                            'volume': volume
                        })
                        
                        total_area += area
                        total_volume += volume
                
                except Exception as e:
                    continue
            
            content.append(f"**Gesamtfläche aus Raumdaten:** {total_area:.2f} m²")
            content.append(f"**Gesamtvolumen aus Raumdaten:** {total_volume:.2f} m³")
            content.append("")
            
            content.append("### Raumübersicht nach Raumtypen")
            content.append("")
            
            for room_type in sorted(room_types.keys()):
                data = room_types[room_type]
                if data['count'] > 0:
                    content.append(f"**{room_type}:**")
                    content.append(f"* Anzahl Räume: {data['count']}")
                    content.append(f"* Gesamtfläche: {data['total_area']:.2f} m²")
                    content.append(f"* Gesamtvolumen: {data['total_volume']:.2f} m³")
                    content.append("")
                    
                    # List first 5 rooms as examples
                    if len(data['rooms']) > 0:
                        content.append("  Beispielräume:")
                        for room in data['rooms'][:5]:
                            content.append(f"  - {room['nr']} {room['name']}: {room['area']:.1f} m²")
                        if len(data['rooms']) > 5:
                            content.append(f"  ... und {len(data['rooms']) - 5} weitere")
                    content.append("")
            
            # Add detailed room-by-room data section
            content.append("## DETAILLIERTE RAUMLISTE")
            content.append("")
            content.append("Nachfolgend die vollständige Liste aller Räume mit technischen Details:")
            content.append("")
            
            # Extract detailed room data
            detailed_rooms = _extract_detailed_room_data(df, col_mapping)
            
            if detailed_rooms:
                for room in detailed_rooms:
                    # Room header
                    room_header = f"**Raum {room.get('raum_nr', 'N/A')}** - {room.get('raum_bezeichnung', 'N/A')}"
                    content.append(room_header)
                    
                    # Basic info
                    if room.get('gebäudeteil'):
                        content.append(f"* Gebäudeteil: {room['gebäudeteil']}")
                    if room.get('geschoss'):
                        content.append(f"* Geschoss: {room['geschoss']}")
                    
                    # Room type with number
                    raumtyp_nr = room.get('raumtyp_nr')
                    raumtyp_bez = room.get('raumtyp_bez')
                    if raumtyp_nr:
                        if raumtyp_bez:
                            content.append(f"* Raumtyp: {raumtyp_bez} ({int(raumtyp_nr)})")
                        else:
                            content.append(f"* Raumtyp Nr.: {int(raumtyp_nr)}")
                    elif raumtyp_bez:
                        content.append(f"* Raumtyp: {raumtyp_bez}")
                    
                    # Dimensions
                    if room.get('fläche'):
                        content.append(f"* Fläche: {room['fläche']:.2f} m²")
                    if room.get('volumen'):
                        content.append(f"* Volumen: {room['volumen']:.2f} m³")
                    if room.get('höhe_mep'):
                        content.append(f"* Höhe MEP-Raum: {room['höhe_mep']:.2f} m")
                    if room.get('lichte_raumhöhe'):
                        content.append(f"* Lichte Raumhöhe: {room['lichte_raumhöhe']:.2f} m")
                    
                    # Technical parameters
                    if room.get('temp_winter'):
                        content.append(f"* Raumtemperatur Winter: {room['temp_winter']}°C")
                    if room.get('temp_sommer'):
                        content.append(f"* Raumtemperatur Sommer: {room['temp_sommer']}°C")
                    if room.get('beheizt'):
                        content.append(f"* Beheizt: {room['beheizt']}")
                    if room.get('mech_lüftung'):
                        content.append(f"* Mechanische Lüftung: {room['mech_lüftung']}")
                    if room.get('luftfeuchte'):
                        content.append(f"* Anforderung Luftfeuchte: {room['luftfeuchte']}")
                    if room.get('akustik'):
                        content.append(f"* Anforderung Akustik: {room['akustik']}")
                    if room.get('sprinkler'):
                        content.append(f"* Sprinklerschutz: {room['sprinkler']}")
                    if room.get('personen_anzahl'):
                        content.append(f"* Anzahl Personen/Gegenstände: {room['personen_anzahl']}")
                    
                    # Ceiling info
                    if room.get('abhangdecke'):
                        content.append(f"* Abhangdecke: {room['abhangdecke']}")
                    if room.get('abhangdecke_typ'):
                        content.append(f"* Typ Abhangdecke: {room['abhangdecke_typ']}")
                    if room.get('lichte_höhe_abhang'):
                        content.append(f"* Lichte Höhe Abhangdecke: {room['lichte_höhe_abhang']:.2f} m")
                    
                    # Comments
                    if room.get('kommentar_1'):
                        content.append(f"* Kommentar (1): {room['kommentar_1']}")
                    if room.get('kommentar_2'):
                        content.append(f"* Kommentar (2): {room['kommentar_2']}")
                    
                    content.append("")
        
        logger.info("Successfully extracted Excel room data")
    except Exception as e:
        logger.error(f"Error extracting Excel data: {e}")
    
    return "\n".join(content)


def extract_project_data(context_directory: Union[str, Path]) -> str:
    """
    Convenience function to extract all project data from context directory
    
    Args:
        context_directory: Path to directory containing project files
        
    Returns:
        Combined project data as text
    """
    extractor = FileExtractor()
    
    # Extract data from all files
    extracted_data = extractor.extract_from_directory(context_directory)
    
    if not extracted_data:
        logger.warning("No data extracted from context directory")
        return ""
    
    # Combine into structured format
    project_data = extractor.combine_extracted_data(extracted_data)
    
    logger.info(f"Successfully extracted data from {len(extracted_data)} files")
    return project_data


if __name__ == "__main__":
    # Test the extractor
    context_dir = Path(__file__).parent.parent.parent / "context"
    if context_dir.exists():
        project_data = extract_project_data(context_dir)
        print(f"Extracted {len(project_data)} characters of project data")
        print("\nFirst 500 characters:")
        print(project_data[:500])
    else:
        print(f"Context directory not found: {context_dir}")
