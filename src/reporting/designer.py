"""Report Designer"""
from datetime import datetime
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.config import REPORTS_DIR
import re


class Designer:
    def __init__(self):
        REPORTS_DIR.mkdir(exist_ok=True)
        self.logo_path = Path(__file__).parent / "bkw_eng_logo.png"
    
    def _filename(self, ext):
        return f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
    
    def _convert_markdown_to_html(self, text):
        """Convert markdown bold (**text**) to HTML <b>text</b>"""
        return re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
    
    def _clean_latex_math(self, text):
        """Convert LaTeX mathematical notation to readable text"""
        replacements = [
            (r'\$([^$]+)\$', lambda m: self._process_math_expression(m.group(1))),
            (r'\$', ''),
        ]
        
        for pattern, replacement in replacements:
            if callable(replacement):
                text = re.sub(pattern, replacement, text)
            else:
                text = re.sub(pattern, replacement, text)
        
        return text
    
    def _process_math_expression(self, expr):
        """Process individual math expressions"""
        expr = re.sub(r'\\text\{([^}]+)\}', r'\1', expr)
        expr = re.sub(r'_{([^}]+)}', r'_\1', expr)
        
        expr = re.sub(r'\^{([^}]+)}', r'^\1', expr)
        
        expr = re.sub(r'm\^3', 'm³', expr)
        expr = re.sub(r'm\^2', 'm²', expr)
        expr = re.sub(r'h\^{-1}', 'h⁻¹', expr)
        expr = re.sub(r'°C', '°C', expr)
        expr = re.sub(r'\^\{\\circ\}C', '°C', expr)
        
        expr = re.sub(r'\s+', ' ', expr).strip()
        
        return expr
    
    def _create_styles(self):
        """Create custom styles for the report"""
        styles = getSampleStyleSheet()
        
        style_definitions = [
            ('CustomTitle', {
                'parent': styles['Heading1'],
                'fontSize': 18,
                'textColor': HexColor('#0066cc'),
                'spaceAfter': 12,
                'spaceBefore': 12,
                'alignment': TA_CENTER,
                'fontName': 'Helvetica-Bold'
            }),
            ('MainHeading', {
                'parent': styles['Heading1'],
                'fontSize': 14,
                'textColor': HexColor('#0066cc'),
                'spaceAfter': 10,
                'spaceBefore': 16,
                'fontName': 'Helvetica-Bold',
                'keepWithNext': True
            }),
            ('SubHeading', {
                'parent': styles['Heading2'],
                'fontSize': 12,
                'textColor': HexColor('#004080'),
                'spaceAfter': 8,
                'spaceBefore': 12,
                'fontName': 'Helvetica-Bold',
                'keepWithNext': True
            }),
            ('BodyText', {
                'parent': styles['Normal'],
                'fontSize': 10,
                'leading': 14,
                'spaceAfter': 6,
                'alignment': TA_JUSTIFY,
                'fontName': 'Helvetica'
            }),
            ('Bullet', {
                'parent': styles['Normal'],
                'fontSize': 10,
                'leading': 14,
                'spaceAfter': 4,
                'leftIndent': 20,
                'bulletIndent': 8,
                'fontName': 'Helvetica'
            })
        ]
        
        for style_name, style_props in style_definitions:
            if style_name not in styles:
                styles.add(ParagraphStyle(name=style_name, **style_props))
        return styles
    
    def _add_header(self, story, doc_title, styles):
        """Add header with logo and title"""
        if self.logo_path.exists():
            logo = Image(str(self.logo_path), width=20*mm, height=20*mm)
            
            title_para = Paragraph(doc_title, styles['CustomTitle'])
            
            header_data = [[title_para, logo]]
            header_table = Table(header_data, colWidths=[165*mm, 20*mm])
            header_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            story.append(header_table)
        else:
            story.append(Paragraph(doc_title, styles['CustomTitle']))
        story.append(Spacer(1, 0.3*inch))
    
    def _parse_content(self, content, styles):
        """Parse content and create story elements"""
        story = []
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            i += 1
            
            if not line:
                story.append(Spacer(1, 0.08*inch))
                continue
            if line.startswith('---'):
                story.append(Spacer(1, 0.1*inch))
                continue
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = self._clean_latex_math(line.lstrip('#').strip())
                text = self._convert_markdown_to_html(text)
                style = styles['MainHeading'] if level <= 2 else styles['SubHeading']
                story.append(Paragraph(text, style))
                continue
            if re.match(r'^A\.\d+\s+KG\s+\d+', line):
                text = self._clean_latex_math(line)
                story.append(Paragraph(self._convert_markdown_to_html(text), styles['MainHeading']))
                continue
            if re.match(r'^\d+\.?\s+[A-ZÄÖÜ]', line) and len(line) < 80:
                text = self._clean_latex_math(line)
                story.append(Paragraph(self._convert_markdown_to_html(text), styles['SubHeading']))
                continue
            if line.startswith('* ') or line.startswith('- '):
                bullet_text = self._clean_latex_math(line[2:].strip())
                story.append(Paragraph(f'• {self._convert_markdown_to_html(bullet_text)}', styles['Bullet']))
                continue
            text = self._clean_latex_math(line)
            story.append(Paragraph(self._convert_markdown_to_html(text), styles['BodyText']))
        return story
    
    def pdf(self, content, doc_title="AI Generated Report"):
        path = REPORTS_DIR / self._filename('.pdf')
        doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=25*mm, rightMargin=25*mm,
                                topMargin=20*mm, bottomMargin=20*mm)
        styles = self._create_styles()
        story = []
        self._add_header(story, doc_title, styles)
        story.append(Paragraph(f"Erstellt: {datetime.now().strftime('%d.%m.%Y')}", styles['BodyText']))
        story.append(Spacer(1, 0.3*inch))
        story.extend(self._parse_content(content, styles))
        doc.build(story)
        return str(path)
    
    def _process_bold_text(self, paragraph, text):
        """Process bold markdown in DOCX"""
        text = self._clean_latex_math(text)
        
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)
    
    def docx(self, content, doc_title="AI Generated Report"):
        path = REPORTS_DIR / self._filename('.docx')
        doc = Document()
        
        if self.logo_path.exists():
            doc.add_picture(str(self.logo_path))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title = doc.add_heading(doc_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para = doc.add_paragraph(f"Erstellt: {datetime.now().strftime('%d.%m.%Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        for line in content.split('\n'):
            line = line.strip()
            if not line or line.startswith('---'):
                doc.add_paragraph()
                continue
            if line.startswith('#'):
                level = min(len(line) - len(line.lstrip('#')), 3)
                heading_text = self._clean_latex_math(line.lstrip('#').strip())
                heading = doc.add_heading(heading_text, level=level)
                if heading.runs:
                    color = RGBColor(0, 102, 204) if level <= 2 else RGBColor(0, 64, 128)
                    heading.runs[0].font.color.rgb = color
                continue
            if re.match(r'^A\.\d+\s+KG\s+\d+', line):
                heading_text = self._clean_latex_math(line)
                heading = doc.add_heading(heading_text, level=1)
                heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
                continue
            if re.match(r'^\d+\.?\s+[A-ZÄÖÜ]', line) and len(line) < 80:
                heading_text = self._clean_latex_math(line)
                heading = doc.add_heading(heading_text, level=2)
                heading.runs[0].font.color.rgb = RGBColor(0, 64, 128)
                continue
            if line.startswith('* ') or line.startswith('- '):
                para = doc.add_paragraph(style='List Bullet')
                self._process_bold_text(para, line[2:].strip())
                continue
            para = doc.add_paragraph()
            self._process_bold_text(para, line)
        
        doc.save(str(path))
        return str(path)
    
    def markdown(self, content):
        path = REPORTS_DIR / self._filename('.md')
        header = f"# Erläuterungsbericht\n\n**Erstellt:** {datetime.now().strftime('%d.%m.%Y')}\n\n---\n\n"
        path.write_text(header + content + "\n", encoding='utf-8')
        return str(path)
